#!/usr/bin/env python3
"""
md_to_docx.py — Convert Markdown text to a styled Word document (.docx).

Standalone utility (like xlsx_to_text.py) that is also importable by run_claude.py.

Features:
  - Headings (#–####) with colored Arial fonts
  - Bold, italic, inline code, strikethrough
  - Bullet and numbered lists
  - Fenced code blocks with Consolas + shading
  - Blockquotes (indented italic)
  - Horizontal rules
  - Markdown tables → Word tables
  - Hyperlinks (colored underlined text)

Requires: python-docx

Usage:
    python md_to_docx.py input.md
    python md_to_docx.py input.md output.docx
    python md_to_docx.py input.txt --title "My Report"
"""

import re
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "Error: python-docx is required. Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


class MarkdownToDocxConverter:
    """Converts Markdown content to a styled Word document."""

    STYLES = {
        "title":      {"font": "Arial",    "size": 28, "bold": True,  "color": "1a365d"},
        "heading1":   {"font": "Arial",    "size": 24, "bold": True,  "color": "2c5282"},
        "heading2":   {"font": "Arial",    "size": 20, "bold": True,  "color": "2b6cb0"},
        "heading3":   {"font": "Arial",    "size": 16, "bold": True,  "color": "3182ce"},
        "heading4":   {"font": "Arial",    "size": 14, "bold": True,  "color": "4299e1"},
        "body":       {"font": "Georgia",  "size": 11,                "color": "2d3748"},
        "code":       {"font": "Consolas", "size": 10,                "color": "c7254e", "bg": "f9f2f4"},
        "code_block": {"font": "Consolas", "size": 10,                "color": "2d3748", "bg": "f7fafc"},
        "blockquote": {"font": "Georgia",  "size": 11, "italic": True,"color": "718096"},
        "link":       {"font": "Georgia",  "size": 11, "underline": True, "color": "3182ce"},
    }

    def __init__(self):
        self.doc = Document()
        self.list_counter = 0
        self._setup_styles()

    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles

        style = styles["Normal"]
        style.font.name = self.STYLES["body"]["font"]
        style.font.size = Pt(self.STYLES["body"]["size"])
        style.font.color.rgb = RGBColor.from_string(self.STYLES["body"]["color"])

        for i, level in enumerate(["heading1", "heading2", "heading3", "heading4"], 1):
            style_name = f"Heading {i}"
            if style_name in styles:
                style = styles[style_name]
                style.font.name = self.STYLES[level]["font"]
                style.font.size = Pt(self.STYLES[level]["size"])
                style.font.bold = self.STYLES[level]["bold"]
                style.font.color.rgb = RGBColor.from_string(self.STYLES[level]["color"])
                style.paragraph_format.space_before = Pt(12)
                style.paragraph_format.space_after = Pt(6)

        if "Title" in styles:
            style = styles["Title"]
            style.font.name = self.STYLES["title"]["font"]
            style.font.size = Pt(self.STYLES["title"]["size"])
            style.font.bold = self.STYLES["title"]["bold"]
            style.font.color.rgb = RGBColor.from_string(self.STYLES["title"]["color"])
            style.paragraph_format.space_after = Pt(18)

    def _add_shading(self, paragraph, color):
        """Add background shading to a paragraph."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        paragraph._p.get_or_add_pPr().append(shading)

    def _parse_inline_formatting(self, text, paragraph):
        """Parse and apply inline formatting (bold, italic, code, links, strikethrough)."""
        patterns = [
            (r'\[([^\]]+)\]\(([^)]+)\)', "link"),
            (r'`([^`]+)`', "code"),
            (r'\*\*\*([^*]+)\*\*\*', "bold_italic"),
            (r'\*\*([^*]+)\*\*', "bold"),
            (r'__([^_]+)__', "bold"),
            (r'\*([^*]+)\*', "italic"),
            (r'_([^_]+)_', "italic"),
            (r'~~([^~]+)~~', "strikethrough"),
        ]

        combined_pattern = "|".join(f"({p[0]})" for p in patterns)

        pos = 0
        for match in re.finditer(combined_pattern, text):
            if match.start() > pos:
                paragraph.add_run(text[pos:match.start()])

            matched_text = match.group(0)

            if matched_text.startswith("[") and "](" in matched_text:
                link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', matched_text)
                if link_match:
                    link_text = link_match.group(1)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor.from_string(self.STYLES["link"]["color"])
                    run.font.underline = True
            elif matched_text.startswith("`"):
                code_text = re.match(r'`([^`]+)`', matched_text).group(1)
                run = paragraph.add_run(code_text)
                run.font.name = self.STYLES["code"]["font"]
                run.font.size = Pt(self.STYLES["code"]["size"])
                run.font.color.rgb = RGBColor.from_string(self.STYLES["code"]["color"])
            elif matched_text.startswith("***"):
                inner_text = re.match(r'\*\*\*([^*]+)\*\*\*', matched_text).group(1)
                run = paragraph.add_run(inner_text)
                run.bold = True
                run.italic = True
            elif matched_text.startswith("**") or matched_text.startswith("__"):
                inner_text = re.match(r'[\*_]{2}([^\*_]+)[\*_]{2}', matched_text).group(1)
                run = paragraph.add_run(inner_text)
                run.bold = True
            elif matched_text.startswith("~~"):
                inner_text = re.match(r'~~([^~]+)~~', matched_text).group(1)
                run = paragraph.add_run(inner_text)
                run.font.strike = True
            elif matched_text.startswith("*") or matched_text.startswith("_"):
                inner_text = re.match(r'[\*_]([^\*_]+)[\*_]', matched_text).group(1)
                run = paragraph.add_run(inner_text)
                run.italic = True

            pos = match.end()

        if pos < len(text):
            paragraph.add_run(text[pos:])

    def _add_paragraph(self, text, style=None):
        """Add a paragraph with inline formatting."""
        p = self.doc.add_paragraph(style=style)
        self._parse_inline_formatting(text, p)
        return p

    def _add_code_block(self, code, language=None):
        """Add a code block with monospace font and shading."""
        for line in code.split("\n"):
            p = self.doc.add_paragraph()
            run = p.add_run(line if line else " ")
            run.font.name = self.STYLES["code_block"]["font"]
            run.font.size = Pt(self.STYLES["code_block"]["size"])
            run.font.color.rgb = RGBColor.from_string(self.STYLES["code_block"]["color"])
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            self._add_shading(p, self.STYLES["code_block"]["bg"])

    def _add_blockquote(self, text):
        """Add a blockquote with special styling."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(self.STYLES["blockquote"]["color"])
        self._add_shading(p, "f0f0f0")
        return p

    def _add_horizontal_rule(self):
        """Add a horizontal rule as a bottom-border paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "cccccc")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_list_item(self, text, ordered=False, level=0):
        """Add a list item (bullet or numbered)."""
        style = "List Number" if ordered else "List Bullet"
        p = self.doc.add_paragraph(style=style)
        p.clear()
        self._parse_inline_formatting(text, p)
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
        return p

    def _add_table(self, rows):
        """Add a Word table from parsed markdown table rows."""
        if not rows:
            return
        # Split each row into cells
        parsed = []
        for row in rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)

        # Skip separator row (e.g. |---|---|)
        data_rows = []
        for row in parsed:
            if all(re.match(r'^[-:]+$', c.strip()) for c in row if c.strip()):
                continue
            data_rows.append(row)

        if not data_rows:
            return

        num_cols = max(len(r) for r in data_rows)
        table = self.doc.add_table(rows=len(data_rows), cols=num_cols, style="Table Grid")

        for r_idx, row_cells in enumerate(data_rows):
            for c_idx, cell_text in enumerate(row_cells):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    self._parse_inline_formatting(cell_text, p)

        # Bold the header row
        if len(data_rows) > 1:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    def convert(self, markdown_text):
        """Convert markdown text to Word document content."""
        lines = markdown_text.split("\n")
        i = 0
        in_code_block = False
        code_block_content = []
        code_block_lang = None

        while i < len(lines):
            line = lines[i]

            # Code block handling
            if line.strip().startswith("```"):
                if in_code_block:
                    self._add_code_block("\n".join(code_block_content), code_block_lang)
                    code_block_content = []
                    code_block_lang = None
                    in_code_block = False
                else:
                    in_code_block = True
                    code_block_lang = line.strip()[3:].strip() or None
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Headings
            if line.startswith("#"):
                match = re.match(r'^(#{1,6})\s+(.+)', line)
                if match:
                    level = len(match.group(1))
                    text = match.group(2)
                    if level == 1:
                        p = self.doc.add_paragraph(style="Title")
                        p.clear()
                        self._parse_inline_formatting(text, p)
                    else:
                        style = f"Heading {min(level, 4)}"
                        p = self.doc.add_paragraph(style=style)
                        p.clear()
                        self._parse_inline_formatting(text, p)
                    i += 1
                    continue

            # Horizontal rule
            if re.match(r'^[-*_]{3,}\s*$', line):
                self._add_horizontal_rule()
                i += 1
                continue

            # Blockquote
            if line.startswith(">"):
                quote_text = line[1:].strip()
                self._add_blockquote(quote_text)
                i += 1
                continue

            # Table (collect consecutive | rows)
            if line.strip().startswith("|") and "|" in line[1:]:
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_rows.append(lines[i])
                    i += 1
                self._add_table(table_rows)
                continue

            # Unordered list
            ul_match = re.match(r'^(\s*)[-*+]\s+(.+)', line)
            if ul_match:
                indent = len(ul_match.group(1))
                level = indent // 2
                text = ul_match.group(2)
                self._add_list_item(text, ordered=False, level=level)
                i += 1
                continue

            # Ordered list
            ol_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
            if ol_match:
                indent = len(ol_match.group(1))
                level = indent // 2
                text = ol_match.group(2)
                self._add_list_item(text, ordered=True, level=level)
                i += 1
                continue

            # Regular paragraph
            self._add_paragraph(line)
            i += 1

        return self.doc

    def save(self, output_path):
        """Save the document to a file."""
        self.doc.save(output_path)


# ── Convenience function (used by run_claude.py) ─────────────────────────────

def markdown_to_docx(text, output_path, title=None):
    """
    Convert markdown text to a .docx file.

    Args:
        text: Markdown-formatted string
        output_path: Destination .docx path
        title: Optional title (unused currently; heading comes from markdown)

    Returns:
        The output_path written to.
    """
    converter = MarkdownToDocxConverter()
    converter.convert(text)
    converter.save(output_path)
    return str(output_path)


# ── CLI entry point ──────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to a styled Word document"
    )
    parser.add_argument("input", help="Input Markdown/text file")
    parser.add_argument("output", nargs="?", help="Output .docx file (default: <input>.docx)")
    parser.add_argument("--title", help="Optional document title")

    args = parser.parse_args()
    input_path = Path(args.input)

    if not input_path.exists():
        print(f"Error: Input file '{input_path}' not found.")
        sys.exit(1)

    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")

    with open(input_path, "r", encoding="utf-8") as f:
        markdown_content = f.read()

    markdown_to_docx(markdown_content, str(output_path), title=args.title)
    print(f"Converted '{input_path}' -> '{output_path}'")


if __name__ == "__main__":
    main()
